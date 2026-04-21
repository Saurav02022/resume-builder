import io
import re
from pdfminer.high_level import extract_text
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdftypes import resolve1
import docx

def parse_pdf_to_text(file_bytes: bytes) -> str:
    """
    Extracts text and hyperlinks from a PDF file byte stream.
    """
    pdf_file = io.BytesIO(file_bytes)
    
    # 1. Extract raw text
    text = extract_text(pdf_file)
    pdf_file.seek(0)
    
    # 2. Extract Hyperlinks from annotations
    links = []
    try:
        parser = PDFParser(pdf_file)
        doc = PDFDocument(parser)
        for page in PDFPage.create_pages(doc):
            if page.annots:
                for annot_ref in resolve1(page.annots):
                    annot = resolve1(annot_ref)
                    if annot.get('Subtype') and annot['Subtype'].name == 'Link':
                        action = resolve1(annot.get('A'))
                        if action:
                            uri = action.get('URI')
                            if uri:
                                try:
                                    uri_str = uri.decode('utf-8')
                                    if uri_str not in links:
                                        links.append(uri_str)
                                except:
                                    pass
    except Exception as e:
        print(f"[PARSER] Warning: Link extraction failed: {str(e)}")

    # 3. Combine text and links for AI visibility
    if links:
        text += "\n\n[SUPPLEMENTAL_LINKS_EXTRACTED_FROM_PDF_METADATA]\n"
        text += "\n".join(links)
        
    return text.strip()

def parse_docx_to_text(file_bytes: bytes) -> str:
    """
    Extracts text from a DOCX file byte stream using python-docx.
    """
    doc_file = io.BytesIO(file_bytes)
    doc = docx.Document(doc_file)
    
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
        
    # Attempt to find URLs in text as well
    urls = re.findall(r'https?://[^\s<>"]+|www\.[^\s<>"]+', '\n'.join(full_text))
    
    result = '\n'.join(full_text)
    if urls:
        result += "\n\n[SUPPLEMENTAL_LINKS_EXTRACTED_FROM_DOCX]\n"
        result += "\n".join(list(set(urls)))
        
    return result.strip()
